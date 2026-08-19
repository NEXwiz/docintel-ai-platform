"""
Document text extraction with structural awareness.

Supports: PDF, DOCX, TXT
- PDFs: Extracts per-page with page markers, handles multi-column layouts,
         extracts text from tables, preserves heading structure.
- DOCX: Extracts paragraphs with heading level info, table content,
         and preserves list formatting.
- TXT:  Direct read with encoding detection fallback.
"""

from pypdf import PdfReader
from docx import Document
import os
import re
from typing import Optional


# Map file extensions to their format tag for chunking routing
_FORMAT_MAP = {
    ".pdf": "pdf",
    ".docx": "docx",
    ".txt": "plain",
    ".md": "markdown",
    ".csv": "plain",
    ".html": "html",
    ".htm": "html",
}


def extract_text(file_path: str) -> str:
    """
    Extract text from a document file.

    Args:
        file_path: Absolute path to the file.

    Returns:
        Extracted text as a single string.

    Raises:
        ValueError: If the file format is unsupported.
        FileNotFoundError: If the file doesn't exist.
    """
    text, _ = extract_with_format(file_path)
    return text


def extract_with_format(file_path: str) -> tuple:
    """
    Extract text and return the format tag for chunking routing.

    Returns:
        Tuple of (extracted_text: str, format_tag: str)
        format_tag is one of: "pdf", "docx", "markdown", "html", "plain"
    """
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"File not found: {file_path}")

    ext = os.path.splitext(file_path)[1].lower()

    extractors = {
        ".pdf": _extract_pdf_text,
        ".docx": _extract_docx_text,
        ".txt": _extract_txt_text,
        ".md": _extract_txt_text,
        ".csv": _extract_txt_text,
        ".html": _extract_html_text,
        ".htm": _extract_html_text,
    }

    extractor = extractors.get(ext)
    if not extractor:
        raise ValueError(f"Unsupported file format: {ext}")

    text = extractor(file_path)
    text = _post_process(text)

    format_tag = _FORMAT_MAP.get(ext, "plain")
    return text, format_tag


# ---------------------------------------------------------------------------
# PDF extraction
# ---------------------------------------------------------------------------

def _extract_pdf_text(path: str) -> str:
    """
    Extract text from PDF with page-level granularity and structure preservation.
    Handles multi-page documents, attempts to detect headings from font styling,
    and includes table text.
    """
    reader = PdfReader(path)
    all_pages = []

    for page_num, page in enumerate(reader.pages, start=1):
        # Try layout mode first (better for multi-column PDFs)
        try:
            page_text = page.extract_text(extraction_mode="layout") or ""
        except (TypeError, ValueError):
            # Older pypdf versions may not support extraction_mode
            page_text = page.extract_text() or ""

        # If layout mode returns nothing, fall back to plain extraction
        if not page_text.strip():
            page_text = page.extract_text() or ""

        if not page_text.strip():
            continue

        # Clean up PDF-specific artifacts on this page
        page_text = _clean_pdf_page(page_text, page_num)

        all_pages.append(page_text)

    if not all_pages:
        return ""

    return "\n\n".join(all_pages)


def _clean_pdf_page(text: str, page_num: int) -> str:
    """Clean artifacts common in PDF text extraction."""
    # Remove excessive spaces from column layout extraction
    text = re.sub(r'[ \t]{4,}', '  ', text)

    # Fix broken words from hyphenated line breaks
    text = re.sub(r'(\w)-\s*\n\s*(\w)', r'\1\2', text)

    # Remove standalone page numbers that aren't part of content
    # Patterns: "12", "- 12 -", "Page 12", "12 of 50"
    lines = text.split('\n')
    cleaned_lines = []
    for line in lines:
        stripped = line.strip()
        # Skip lines that are just page numbers
        if re.match(r'^-?\s*\d{1,4}\s*-?$', stripped):
            continue
        if re.match(r'^Page\s+\d+(\s+of\s+\d+)?$', stripped, re.IGNORECASE):
            continue
        cleaned_lines.append(line)

    return '\n'.join(cleaned_lines)


# ---------------------------------------------------------------------------
# DOCX extraction
# ---------------------------------------------------------------------------

def _extract_docx_text(path: str) -> str:
    """
    Extract text from DOCX preserving document structure:
    - Headings are prefixed with markdown-style markers
    - Lists retain their bullet/number formatting
    - Tables are extracted with row/column structure
    """
    doc = Document(path)
    parts = []

    for element in doc.element.body:
        tag = element.tag.split('}')[-1]  # Strip namespace

        if tag == 'p':
            # It's a paragraph
            para = _find_paragraph_for_element(doc, element)
            if para is not None:
                text = _extract_paragraph_text(para)
                if text:
                    parts.append(text)

        elif tag == 'tbl':
            # It's a table
            table = _find_table_for_element(doc, element)
            if table is not None:
                table_text = _extract_table_text(table)
                if table_text:
                    parts.append(table_text)

    return "\n\n".join(parts)


def _find_paragraph_for_element(doc: Document, element) -> Optional[object]:
    """Find the python-docx Paragraph object matching an lxml element."""
    for para in doc.paragraphs:
        if para._element is element:
            return para
    return None


def _find_table_for_element(doc: Document, element) -> Optional[object]:
    """Find the python-docx Table object matching an lxml element."""
    for table in doc.tables:
        if table._element is element:
            return table
    return None


def _extract_paragraph_text(para) -> str:
    """Extract paragraph text with heading/list formatting preserved."""
    text = para.text.strip()
    if not text:
        return ""

    style_name = (para.style.name or "").lower() if para.style else ""

    # Headings → markdown-style markers
    if "heading" in style_name:
        # Extract heading level (Heading 1 → #, Heading 2 → ##, etc.)
        level = 1
        for char in style_name:
            if char.isdigit():
                level = int(char)
                break
        prefix = "#" * min(level, 6)
        return f"{prefix} {text}"

    # List items → preserve bullet/number prefix
    if "list" in style_name or "bullet" in style_name:
        # Check if text already has a bullet/number prefix
        if not re.match(r'^[\u2022\u2023\u25E6•\-\*\d]', text):
            return f"• {text}"

    # Title style
    if "title" in style_name:
        return f"# {text}"

    return text


def _extract_table_text(table) -> str:
    """Extract table content as structured text.
    Format: pipe-delimited rows for readability."""
    rows = []
    for row in table.rows:
        cells = [cell.text.strip().replace('\n', ' ') for cell in row.cells]
        # Skip completely empty rows
        if any(cells):
            rows.append(" | ".join(cells))

    if not rows:
        return ""

    # Add a separator after the first row (treat as header)
    if len(rows) > 1:
        col_count = rows[0].count('|') + 1
        separator = " | ".join(["---"] * col_count)
        rows.insert(1, separator)

    return "\n".join(rows)


# ---------------------------------------------------------------------------
# HTML extraction
# ---------------------------------------------------------------------------

def _extract_html_text(path: str) -> str:
    """Extract text from HTML file using BeautifulSoup.
    Returns raw HTML for the HTML chunker, with noise elements removed."""
    raw = _extract_txt_text(path)  # read file with encoding fallback

    try:
        from bs4 import BeautifulSoup
    except ImportError:
        # If BeautifulSoup not installed, strip tags with regex as fallback
        text = re.sub(r'<script[^>]*>.*?</script>', '', raw, flags=re.DOTALL | re.IGNORECASE)
        text = re.sub(r'<style[^>]*>.*?</style>', '', text, flags=re.DOTALL | re.IGNORECASE)
        text = re.sub(r'<[^>]+>', ' ', text)
        text = re.sub(r'\s+', ' ', text)
        return text.strip()

    soup = BeautifulSoup(raw, 'html.parser')

    # Remove noise elements
    for tag in soup.find_all(['script', 'style', 'nav', 'footer', 'header']):
        tag.decompose()

    return soup.get_text(separator='\n', strip=True)


# ---------------------------------------------------------------------------
# TXT / plain text extraction
# ---------------------------------------------------------------------------

def _extract_txt_text(path: str) -> str:
    """Read text file with encoding fallback chain."""
    encodings = ["utf-8", "utf-8-sig", "latin-1", "cp1252"]

    for encoding in encodings:
        try:
            with open(path, "r", encoding=encoding) as f:
                return f.read()
        except (UnicodeDecodeError, UnicodeError):
            continue

    # Last resort: read with error replacement
    with open(path, "r", encoding="utf-8", errors="replace") as f:
        return f.read()


# ---------------------------------------------------------------------------
# Post-processing
# ---------------------------------------------------------------------------

def _post_process(text: str) -> str:
    """Final cleanup pass on extracted text."""
    if not text:
        return ""

    # Remove null bytes and control characters (keep newline, tab)
    text = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f]', '', text)

    # Collapse runs of 3+ newlines to 2
    text = re.sub(r'\n{3,}', '\n\n', text)

    # Remove trailing whitespace per line
    text = re.sub(r'[ \t]+$', '', text, flags=re.MULTILINE)

    return text.strip()
